#Hàm thêm một đoạn văn vào sau một đoạn văn chỉ định
def insert_paragraph_after(paragraph, text):
    import docx
    from docx.text.paragraph import Paragraph
    from docx.oxml.xmlchemy import OxmlElement
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Pt, Inches
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.add_run(text)
    new_para.paragraph_format.first_line_indent = Inches(.5)
    new_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in new_para.runs:
        run.font.size = Pt(14)
    return new_para
#Hàm xóa đi một đoạn trong word
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

#Tô đậm một đoạn văn bản tượng trong đoạn văn
def bold_selection(para, text, size):
    from docx.shared import Pt
    text = text.removeprefix(" ")
    text = text.removesuffix(" ")
    content = para.text
    para.text = ""
    content = content.split(text)
    count = 1
    for char in content:
        para.add_run(char).font.bold = False
        if count < len(content):
            para.add_run(text).font.bold = True
        count +=1 
    for run in para.runs:
        run.font.size = Pt(size)